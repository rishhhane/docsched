import os
import docx
import re
import datetime
import calendar
from app import create_app
from app.models import db, Doctor, Leave, Schedule, ScheduleMeta

months_map = {
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
}

def clean_doctor_name(name):
    name = name.strip()
    if not name or name.lower() in ["", "none", "off", "leave", "-", "nil"]:
        return None
    # Split on slashes or "or" / "and" / " / "
    if "/" in name:
        name = name.split("/")[0].strip()
    if "\\" in name:
        name = name.split("\\")[0].strip()
    if " or " in name.lower():
        name = re.split(r'\s+or\s+', name, flags=re.IGNORECASE)[0].strip()
    # Typo overrides
    if name.lower() == "pavanaiko":
        return "Pavan"
    if name.lower() == "tirthankar":
        return "Thirthankar"
    if name.lower() == "utpala":
        return "Uthpala"
    return name

def parse_filename(filename):
    lower = filename.lower()
    m = None
    for name, num in months_map.items():
        if name in lower:
            m = num
            break
    if not m:
        return None
        
    year = 2026
    year_match = re.search(r'20\d{2}', filename)
    if year_match:
        year = int(year_match.group(0))
        
    is_final = "final" in lower or "modified" in lower
    return {"month": m, "year": year, "is_final": is_final, "filename": filename}

def parse_docx(file_path):
    doc = docx.Document(file_path)
    if not doc.tables:
        return None
        
    table = doc.tables[0]
    rows_data = []
    
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        
        # De-duplicate adjacent cells due to merged cells
        unique_cells = []
        for text in cells_text:
            if not unique_cells or unique_cells[-1] != text:
                unique_cells.append(text)
                
        if len(unique_cells) < 2:
            continue
            
        date_str = unique_cells[0]
        if not re.match(r'^\d+', date_str):
            continue
            
        day_match = re.match(r'^(\d+)', date_str)
        day_num = int(day_match.group(1))
        
        # Pad unique_cells to length 7 to prevent IndexErrors
        while len(unique_cells) < 7:
            unique_cells.append('')
            
        m1 = clean_doctor_name(unique_cells[1])
        m2 = clean_doctor_name(unique_cells[2])
        m3 = clean_doctor_name(unique_cells[3])
        e1 = clean_doctor_name(unique_cells[4])
        e2 = clean_doctor_name(unique_cells[5])
        e3 = clean_doctor_name(unique_cells[6])
        
        rows_data.append({
            "day": day_num,
            "morning": [m1, m2, m3],
            "evening": [e1, e2, e3]
        })
        
    return rows_data

def run_import():
    app = create_app()
    with app.app_context():
        parent_dir = ".."
        files = [f for f in os.listdir(parent_dir) if f.endswith('.docx')]
        
        parsed_files = []
        for f in files:
            meta = parse_filename(f)
            if meta:
                parsed_files.append(meta)
                
        # Group by month/year and pick the preferred one (final/modified)
        grouped = {}
        for pf in parsed_files:
            key = (pf["month"], pf["year"])
            if key not in grouped:
                grouped[key] = pf
            else:
                if pf["is_final"] and not grouped[key]["is_final"]:
                    grouped[key] = pf

        print(f"Clearing existing database tables...")
        # Clear database to perform a clean import of the year's data
        db.session.query(Schedule).delete()
        db.session.query(ScheduleMeta).delete()
        db.session.query(Leave).delete()
        db.session.query(Doctor).delete()
        db.session.commit()

        # Step 1: Pass through all schedules to extract doctors and count slot positions
        doctor_appearances = {}
        monthly_schedules = {}
        
        for key, pf in sorted(grouped.items()):
            file_path = os.path.join(parent_dir, pf["filename"])
            schedule_data = parse_docx(file_path)
            if not schedule_data:
                continue
                
            monthly_schedules[key] = schedule_data
            
            for day in schedule_data:
                for idx, doc_name in enumerate(day["morning"]):
                    if doc_name:
                        doctor_appearances.setdefault(doc_name, [0, 0, 0])[idx] += 1
                for idx, doc_name in enumerate(day["evening"]):
                    if doc_name:
                        doctor_appearances.setdefault(doc_name, [0, 0, 0])[idx] += 1

        # Step 2: Register doctors with calculated priorities
        print(f"Importing {len(doctor_appearances)} doctors...")
        doctor_objs = {}
        for name, counts in sorted(doctor_appearances.items()):
            # Priority determined by slot where doctor appears most frequently
            highest_slot_idx = counts.index(max(counts))
            priority = highest_slot_idx + 1
            
            doc = Doctor(name=name, priority=priority)
            db.session.add(doc)
            db.session.flush() # Flush to generate ID
            doctor_objs[name.lower()] = doc
            print(f"  Added {name} (Priority {priority})")
        db.session.commit()

        # Step 3: Insert schedules
        print("Importing schedule entries...")
        for (month, year), schedule_data in sorted(monthly_schedules.items()):
            print(f"  Importing {month}/{year}...")
            
            # Create metadata entry
            meta = ScheduleMeta(month=month, year=year)
            db.session.add(meta)
            
            for day_data in schedule_data:
                day_num = day_data["day"]
                # Validate date safety
                last_day_of_month = calendar.monthrange(year, month)[1]
                if day_num > last_day_of_month:
                    print(f"    Skipping invalid day {day_num} for month {month}/{year}")
                    continue
                    
                date_obj = datetime.date(year, month, day_num)
                
                # Morning Shift
                m_docs = [doctor_objs.get(name.lower()) if name else None for name in day_data["morning"]]
                sched_morning = Schedule(
                    date=date_obj,
                    shift='morning',
                    doctor_1_id=m_docs[0].id if m_docs[0] else None,
                    doctor_2_id=m_docs[1].id if m_docs[1] else None,
                    doctor_3_id=m_docs[2].id if m_docs[2] else None
                )
                db.session.add(sched_morning)
                
                # Evening Shift
                e_docs = [doctor_objs.get(name.lower()) if name else None for name in day_data["evening"]]
                sched_evening = Schedule(
                    date=date_obj,
                    shift='evening',
                    doctor_1_id=e_docs[0].id if e_docs[0] else None,
                    doctor_2_id=e_docs[1].id if e_docs[1] else None,
                    doctor_3_id=e_docs[2].id if e_docs[2] else None
                )
                db.session.add(sched_evening)
                
            db.session.commit()
            
        print("Successfully imported all sample schedules!")

if __name__ == '__main__':
    run_import()
