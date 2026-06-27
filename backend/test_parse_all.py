import os
import docx
import re

months_map = {
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
}

def parse_filename(filename):
    lower = filename.lower()
    m = None
    for name, num in months_map.items():
        if name in lower:
            m = num
            break
    if not m:
        return None
        
    year = 2026
    year_match = re.search(r'20\d{2}', filename)
    if year_match:
        year = int(year_match.group(0))
        
    is_final = "final" in lower or "modified" in lower
    return {"month": m, "year": year, "is_final": is_final, "filename": filename}

def parse_docx(file_path):
    doc = docx.Document(file_path)
    if not doc.tables:
        return None
        
    table = doc.tables[0]
    rows_data = []
    
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        
        # De-duplicate adjacent cells due to merged cells
        unique_cells = []
        for text in cells_text:
            if not unique_cells or unique_cells[-1] != text:
                unique_cells.append(text)
                
        # We expect date, morning 1-3, evening 1-3 (columns: 7 elements)
        # If it matches a header row or is too short, skip
        if len(unique_cells) < 2:
            continue
            
        date_str = unique_cells[0]
        # Check if first element is a date (starts with digits, or matches \d+/\d+)
        if not re.match(r'^\d+', date_str):
            continue
            
        # Extract day number
        day_match = re.match(r'^(\d+)', date_str)
        day_num = int(day_match.group(1))
        
        # Extract doctor slots
        # Pad unique_cells to length 7 to prevent IndexErrors
        while len(unique_cells) < 7:
            unique_cells.append('')
            
        m1, m2, m3 = unique_cells[1], unique_cells[2], unique_cells[3]
        e1, e2, e3 = unique_cells[4], unique_cells[5], unique_cells[6]
        
        rows_data.append({
            "day": day_num,
            "morning": [m1, m2, m3],
            "evening": [e1, e2, e3]
        })
        
    return rows_data

def run():
    parent_dir = ".."
    files = [f for f in os.listdir(parent_dir) if f.endswith('.docx')]
    
    parsed_files = []
    for f in files:
        meta = parse_filename(f)
        if meta:
            parsed_files.append(meta)
            
    # Group by month/year and pick the preferred one (final/modified)
    grouped = {}
    for pf in parsed_files:
        key = (pf["month"], pf["year"])
        if key not in grouped:
            grouped[key] = pf
        else:
            # If current is final/modified, prefer it
            if pf["is_final"] and not grouped[key]["is_final"]:
                grouped[key] = pf
                
    print(f"Discovered rotas for {len(grouped)} unique months:")
    doctor_appearances = {}
    
    for (m, y), pf in sorted(grouped.items()):
        file_path = os.path.join(parent_dir, pf["filename"])
        print(f"  Month: {m}/{y} -> {pf['filename']} (Final: {pf['is_final']})")
        
        schedule_data = parse_docx(file_path)
        if not schedule_data:
            print("    Failed to parse table")
            continue
            
        print(f"    Parsed {len(schedule_data)} days")
        
        for day in schedule_data:
            # Track doctor appearances by slot (to determine priorities)
            for idx, doc_name in enumerate(day["morning"]):
                if doc_name and doc_name.lower() not in ["", "none", "off"]:
                    # Clean up slash splits like "Arun / Vinay"
                    names = [n.strip() for n in re.split(r'[/]', doc_name)]
                    for name in names:
                        if name:
                            # We record slot index (0 = Slot 1, 1 = Slot 2, 2 = Slot 3)
                            doctor_appearances.setdefault(name, [0, 0, 0])[idx] += 1
            for idx, doc_name in enumerate(day["evening"]):
                if doc_name and doc_name.lower() not in ["", "none", "off"]:
                    names = [n.strip() for n in re.split(r'[/]', doc_name)]
                    for name in names:
                        if name:
                            doctor_appearances.setdefault(name, [0, 0, 0])[idx] += 1

    print("\n--- Discovered Doctors & Priorities ---")
    for doc, counts in sorted(doctor_appearances.items()):
        # Determine priority by highest count in slots
        highest_slot_idx = counts.index(max(counts))
        priority = highest_slot_idx + 1
        print(f"  {doc}: Priority {priority} (Counts in Slot 1/2/3: {counts})")

if __name__ == '__main__':
    run()
