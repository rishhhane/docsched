import docx
import os

def inspect(file_path):
    print(f"Inspecting: {file_path}")
    doc = docx.Document(file_path)
    
    # Print paragraphs
    print("--- Paragraphs ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
            
    # Print tables
    print("\n--- Tables ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table {t_idx} (Rows: {len(table.rows)}):")
        # Print first few rows
        for r_idx in range(min(len(table.rows), 10)):
            row = table.rows[r_idx]
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # De-duplicate adjacent identical cells (due to merged cells)
            merged_row_text = []
            for text in row_text:
                if not merged_row_text or merged_row_text[-1] != text:
                    merged_row_text.append(text)
            print(f"  Row {r_idx}: {merged_row_text}")

if __name__ == '__main__':
    parent_dir = ".."
    files = [f for f in os.listdir(parent_dir) if f.endswith('.docx')]
    if files:
        # Sort files so we inspect one of them
        files.sort()
        inspect(os.path.join(parent_dir, files[0]))
    else:
        print("No docx files found in parent directory")
