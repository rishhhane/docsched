from flask import Blueprint, send_file, request, jsonify, session
import io
import datetime
import calendar
from app.models import db, Schedule, Doctor, Leave
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

export_bp = Blueprint('export', __name__)


def add_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor('#6b7280'))
    # Landscape page size is 792 x 612. Margin is 30.
    page_num = canvas.getPageNumber()
    canvas.drawString(30, 20, "MedScheduler Portal — Confidential")
    canvas.drawRightString(762, 20, f"Page {page_num}")
    canvas.restoreState()

@export_bp.route('/api/export/pdf', methods=['GET'])
def export_pdf():
    month = request.args.get('month', type=int)
    year = request.args.get('year', type=int)
    
    if not month or not year:
        now = datetime.datetime.now()
        month = month or now.month
        year = year or now.year
        
    start_date = datetime.date(year, month, 1)
    last_day = calendar.monthrange(year, month)[1]
    end_date = datetime.date(year, month, last_day)
    
    schedules = Schedule.query.filter(Schedule.date >= start_date, Schedule.date <= end_date, Schedule.user_id == session['user_id']).order_by(Schedule.date, Schedule.shift).all()
    doctors = Doctor.query.filter_by(user_id=session['user_id']).order_by(Doctor.priority, Doctor.name).all()
    leaves = Leave.query.filter(Leave.leave_date >= start_date, Leave.leave_date <= end_date, Leave.user_id == session['user_id']).all()
    
    # Process stats
    leaves_by_doctor = {}
    for l in leaves:
        leaves_by_doctor.setdefault(l.doctor_id, []).append(l.leave_date.strftime("%d"))
        
    stats = {
        doc.id: {
            'name': doc.name,
            'priority': doc.priority,
            'morning': 0,
            'evening': 0,
            'total': 0,
            'leaves': sorted(leaves_by_doctor.get(doc.id, []), key=int)
        }
        for doc in doctors
    }
    
    for s in schedules:
        for doc_id, shift_type in [(s.doctor_1_id, s.shift), (s.doctor_2_id, s.shift), (s.doctor_3_id, s.shift)]:
            if doc_id and doc_id in stats:
                stats[doc_id]['total'] += 1
                if shift_type == 'morning':
                    stats[doc_id]['morning'] += 1
                elif shift_type == 'evening':
                    stats[doc_id]['evening'] += 1

    # Group schedules by date
    schedules_by_date = {}
    for s in schedules:
        schedules_by_date.setdefault(s.date, {})[s.shift] = s

    # PDF construction in Landscape mode
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(letter),
        rightMargin=25, leftMargin=25, topMargin=25, bottomMargin=25
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=5
    )
    section_style = ParagraphStyle(
        'DocSection',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        textColor=colors.HexColor('#334155'),
        spaceBefore=10,
        spaceAfter=5
    )
    
    story = []
    
    # Title
    month_name = calendar.month_name[month]
    story.append(Paragraph(f"Doctor Schedule Management Portal — {month_name} {year}", title_style))
    story.append(Spacer(1, 5))
    
    # 1. Main Unified Schedule Table (Side-by-Side Shifts)
    # Row 0: Spanned headers for Morning and Evening
    # Row 1: Sub-headers for Priority slots
    table_data = [
        ["Date", "Day", "Day Shift", "", "", "Evening Shift", "", ""],
        ["", "", "Slot 1", "Slot 2", "Slot 3", "Slot 1", "Slot 2", "Slot 3"]
    ]
    
    # Add each calendar day's assignments
    current_day = start_date
    row_idx = 2
    sunday_rows = []
    while current_day <= end_date:
        if current_day.weekday() == 6:  # Sunday
            sunday_rows.append(row_idx)
            
        day_str = current_day.strftime("%d")
        day_name = current_day.strftime("%a")
        
        day_shifts = schedules_by_date.get(current_day, {})
        m_shift = day_shifts.get('morning')
        e_shift = day_shifts.get('evening')
        
        m_doc1 = m_shift.doctor_1.name if (m_shift and m_shift.doctor_1) else "None"
        m_doc2 = m_shift.doctor_2.name if (m_shift and m_shift.doctor_2) else "None"
        m_doc3 = m_shift.doctor_3.name if (m_shift and m_shift.doctor_3) else "None"
        
        # Check for duplicates in morning shift
        m_docs = [m_doc1, m_doc2, m_doc3]
        m_duplicates = {x for x in m_docs if x != "None" and m_docs.count(x) > 1}
        
        if m_doc1 in m_duplicates:
            m_doc1 = f"{m_doc1} (Duplicate)"
        if m_doc2 in m_duplicates:
            m_doc2 = f"{m_doc2} (Duplicate)"
        if m_doc3 in m_duplicates:
            m_doc3 = f"{m_doc3} (Duplicate)"
            
        e_doc1 = e_shift.doctor_1.name if (e_shift and e_shift.doctor_1) else "None"
        e_doc2 = e_shift.doctor_2.name if (e_shift and e_shift.doctor_2) else "None"
        e_doc3 = e_shift.doctor_3.name if (e_shift and e_shift.doctor_3) else "None"
        
        # Check for duplicates in evening shift
        e_docs = [e_doc1, e_doc2, e_doc3]
        e_duplicates = {x for x in e_docs if x != "None" and e_docs.count(x) > 1}
        
        if e_doc1 in e_duplicates:
            e_doc1 = f"{e_doc1} (Duplicate)"
        if e_doc2 in e_duplicates:
            e_doc2 = f"{e_doc2} (Duplicate)"
        if e_doc3 in e_duplicates:
            e_doc3 = f"{e_doc3} (Duplicate)"
        
        table_data.append([
            day_str,
            day_name,
            m_doc1,
            m_doc2,
            m_doc3,
            e_doc1,
            e_doc2,
            e_doc3
        ])
        row_idx += 1
        current_day += datetime.timedelta(days=1)

    # Calculate maximum font size for Page 1 table to fit in 530 points
    num_data_rows = last_day
    fs_page1 = (530 - 12 - 4 * num_data_rows) / (num_data_rows + 2)
    fs_page1 = int(fs_page1 * 2) / 2.0
    fs_page1 = min(fs_page1, 12.0)

    header_height = fs_page1 + 6.0
    row_height = fs_page1 + 4.0
    row_heights = [header_height, header_height] + [row_height] * (len(table_data) - 2)
        
    t_style = [
        ('SPAN', (0,0), (0,1)), # Span Date header
        ('SPAN', (1,0), (1,1)), # Span Day header
        ('SPAN', (2,0), (4,0)), # Span Morning Shift header
        ('SPAN', (5,0), (7,0)), # Span Evening Shift header
        
        ('BACKGROUND', (0,0), (1,1), colors.HexColor('#f1f5f9')), # light gray header corner
        ('BACKGROUND', (2,0), (4,0), colors.HexColor('#e2e8f0')), # Morning Header (slate-200)
        ('BACKGROUND', (5,0), (7,0), colors.HexColor('#cbd5e1')), # Evening Header (slate-300)
        ('BACKGROUND', (2,1), (4,1), colors.HexColor('#f8fafc')), # Morning Subheaders (slate-50)
        ('BACKGROUND', (5,1), (7,1), colors.HexColor('#f1f5f9')), # Evening Subheaders (slate-100)
        
        ('TEXTCOLOR', (0,0), (-1,1), colors.HexColor('#1f2937')), # dark text for headers
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('FONTNAME', (0,0), (-1,1), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), fs_page1),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2.0),
        ('TOPPADDING', (0,0), (-1,-1), 2.0),
    ]
    
    # Alternating data row backgrounds (clean light slate/white look)
    for r in range(2, len(table_data)):
        if r % 2 == 0:
            t_style.append(('BACKGROUND', (0, r), (-1, r), colors.HexColor('#f8fafc')))
            
    # Bold Sunday rows
    for r in sunday_rows:
        t_style.append(('FONTNAME', (0, r), (-1, r), 'Helvetica-Bold'))
                
    # colWidths total = 35*2 + 110*6 = 70 + 660 = 730 points. Page width is 792. Printable area is 732.
    t_schedule = Table(table_data, colWidths=[35, 35, 110, 110, 110, 110, 110, 110], rowHeights=row_heights)
    t_schedule.setStyle(TableStyle(t_style))
    story.append(t_schedule)
    
    # 2. Page Break for Stats and Rankings
    story.append(PageBreak())
    
    # Page 2 Title
    story.append(Paragraph("Doctor Workload Statistics & Equity Report", title_style))
    story.append(Spacer(1, 10))
    
    # Workload Stats Table data (340 points wide)
    stats_data = [["Doctor Name", "Prio", "Morn", "Eve", "Total"]]
    for doc_id, item in stats.items():
        stats_data.append([
            item['name'],
            f"P{item['priority']}",
            str(item['morning']),
            str(item['evening']),
            str(item['total'])
        ])
    # Calculate maximum font size for Page 2 tables to fit in 530 points
    num_doctors = len(doctors)
    if num_doctors > 0:
        fs_page2 = (530 - 6 - 4 * num_doctors) / (num_doctors + 1)
        fs_page2 = int(fs_page2 * 2) / 2.0
        fs_page2 = min(fs_page2, 12.0)
    else:
        fs_page2 = 12.0

    stats_header_height = fs_page2 + 6.0
    stats_row_height = fs_page2 + 4.0
    stats_row_heights = [stats_header_height] + [stats_row_height] * (len(stats_data) - 1)

    stats_style = [
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#475569')), # slate-600
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('ALIGN', (0,1), (0,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), fs_page2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2.0),
        ('TOPPADDING', (0,0), (-1,-1), 2.0),
    ]
    
    # Alternating row colors for stats table
    for row_idx in range(1, len(stats_data)):
        if row_idx % 2 == 0:
            stats_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), colors.HexColor('#f8fafc')))
            
    t_stats = Table(stats_data, colWidths=[120, 40, 60, 60, 60], rowHeights=stats_row_heights)
    t_stats.setStyle(TableStyle(stats_style))
    
    leaves_style = ParagraphStyle(
        'LeavesStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=fs_page2,
        leading=fs_page2 + 2.0,
        alignment=1,  # Center alignment
        textColor=colors.HexColor('#1f2937')
    )
    
    # Rankings and Leaves Table data (340 points wide)
    ranked_stats = sorted(stats.values(), key=lambda x: (x['total'], x['evening'], x['name']))
    rankings_data = [["Rank", "Doctor Name", "Leave Dates", "Equity Status"]]
    
    for i, item in enumerate(ranked_stats):
        avg_shifts = sum(x['total'] for x in ranked_stats) / len(ranked_stats) if ranked_stats else 0
        diff = item['total'] - avg_shifts
        if diff > 3:
            status = "Overworked"
        elif diff < -3:
            status = "Underutilized"
        else:
            status = "Balanced"
            
        leaves_str = ", ".join(item['leaves']) if item['leaves'] else "None"
        rankings_data.append([
            str(i+1),
            item['name'],
            Paragraph(leaves_str, leaves_style),
            status
        ])
        
    rankings_style = [
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#475569')), # slate-600
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('ALIGN', (1,1), (1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), fs_page2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2.0),
        ('TOPPADDING', (0,0), (-1,-1), 2.0),
    ]
    
    # Alternating row colors for rankings table
    for row_idx in range(1, len(rankings_data)):
        if row_idx % 2 == 0:
            rankings_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), colors.HexColor('#f8fafc')))
            
    t_rankings = Table(rankings_data, colWidths=[40, 110, 110, 80])
    t_rankings.setStyle(TableStyle(rankings_style))
    
    # Place Workload Stats and Rankings side-by-side using an outer Table wrapper
    # Total width: 340 + 50 (gap) + 340 = 730 points
    stats_and_rankings_table = Table([[t_stats, Spacer(1, 1), t_rankings]], colWidths=[340, 50, 340])
    stats_and_rankings_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
    ]))
    
    story.append(stats_and_rankings_table)
    
    # Build PDF and register footer template
    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    buffer.seek(0)
    
    filename = f"schedule_{year}_{month:02d}.pdf"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )

@export_bp.route('/api/export/docx', methods=['GET'])
def export_docx():
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    month = request.args.get('month', type=int)
    year = request.args.get('year', type=int)
    
    if not month or not year:
        now = datetime.datetime.now()
        month = month or now.month
        year = year or now.year
        
    start_date = datetime.date(year, month, 1)
    last_day = calendar.monthrange(year, month)[1]
    end_date = datetime.date(year, month, last_day)
    
    schedules = Schedule.query.filter(Schedule.date >= start_date, Schedule.date <= end_date, Schedule.user_id == session['user_id']).order_by(Schedule.date, Schedule.shift).all()
    doctors = Doctor.query.filter_by(user_id=session['user_id']).order_by(Doctor.priority, Doctor.name).all()
    leaves = Leave.query.filter(Leave.leave_date >= start_date, Leave.leave_date <= end_date, Leave.user_id == session['user_id']).all()
    
    # Process stats
    leaves_by_doctor = {}
    for l in leaves:
        leaves_by_doctor.setdefault(l.doctor_id, []).append(l.leave_date.strftime("%d"))
        
    stats = {
        doc.id: {
            'name': doc.name,
            'priority': doc.priority,
            'morning': 0,
            'evening': 0,
            'total': 0,
            'leaves': sorted(leaves_by_doctor.get(doc.id, []), key=int)
        }
        for doc in doctors
    }
    
    for s in schedules:
        for doc_id, shift_type in [(s.doctor_1_id, s.shift), (s.doctor_2_id, s.shift), (s.doctor_3_id, s.shift)]:
            if doc_id and doc_id in stats:
                stats[doc_id]['total'] += 1
                if shift_type == 'morning':
                    stats[doc_id]['morning'] += 1
                elif shift_type == 'evening':
                    stats[doc_id]['evening'] += 1

    # Group schedules by date
    schedules_by_date = {}
    for s in schedules:
        schedules_by_date.setdefault(s.date, {})[s.shift] = s

    # Create document
    doc = docx.Document()
    
    # Page setup - Landscape with tight margins for 1-page fit
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    
    # Helper to set cell background color
    def set_cell_background(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    # Helper to set table borders
    def set_table_borders(table, color="cbd5e1", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    # Helper to clear table borders
    def clear_table_borders(table):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:bottom w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideH w:val="none"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    # Helper to set cell padding
    def set_cell_margins(cell, top=20, bottom=20, left=50, right=50):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="{top}" w:type="dxa"/>'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'  <w:left w:w="{left}" w:type="dxa"/>'
            f'  <w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    # Helper to format text inside cells
    def format_cell_text(cell, text, font_name="Arial", size_pt=7.5, bold=False, color_hex="1f2937", alignment=WD_ALIGN_PARAGRAPH.CENTER):
        p = cell.paragraphs[0]
        p.text = ""
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if text:
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run.bold = bold
            if color_hex:
                hex_str = color_hex.lstrip('#')
                run.font.color.rgb = RGBColor(*(int(hex_str[i:i+2], 16) for i in (0, 2, 4)))

    month_name = calendar.month_name[month]
    
    # 1. Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    p_title_run = p_title.add_run(f"Doctor Schedule Management Portal — {month_name} {year}")
    p_title_run.font.name = "Arial"
    p_title_run.font.size = Pt(12)
    p_title_run.bold = True
    p_title_run.font.color.rgb = RGBColor(30, 41, 59) # slate-800
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(3)
    p_sub_run = p_sub.add_run("Confidential — Generated Rota Schedule")
    p_sub_run.font.name = "Arial"
    p_sub_run.font.size = Pt(8)
    p_sub_run.font.color.rgb = RGBColor(107, 114, 128) # slate-500

    # Small Heading
    p_head = doc.add_paragraph()
    p_head.paragraph_format.space_before = Pt(0)
    p_head.paragraph_format.space_after = Pt(2)
    p_head_run = p_head.add_run("Monthly Rota Schedule")
    p_head_run.font.name = "Arial"
    p_head_run.font.size = Pt(9.5)
    p_head_run.bold = True
    p_head_run.font.color.rgb = RGBColor(51, 65, 85)

    # Main Table setup (no spacing between cells, clean borders)
    # 8 columns: Date, Day, Morning P1, P2, P3, Evening P1, P2, P3
    table = doc.add_table(rows=2, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="cbd5e1", sz="4")
    
    # Col widths (inches)
    col_widths = [0.45, 0.45, 1.55, 1.55, 1.55, 1.55, 1.55, 1.55]
    
    # Header cells formatting
    cells_0 = table.rows[0].cells
    cells_1 = table.rows[1].cells
    
    # Merge Date & Day vertically
    cells_0[0].merge(cells_1[0])
    cells_0[1].merge(cells_1[1])
    
    # Merge Morning & Evening horizontally
    cells_0[2].merge(cells_0[3]).merge(cells_0[4])
    cells_0[5].merge(cells_0[6]).merge(cells_0[7])
    
    # Set headers text
    format_cell_text(cells_0[0], "Date", size_pt=8.5, bold=True, color_hex="1f2937")
    format_cell_text(cells_0[1], "Day", size_pt=8.5, bold=True, color_hex="1f2937")
    format_cell_text(cells_0[2], "Day Shift", size_pt=8.5, bold=True, color_hex="1f2937")
    format_cell_text(cells_0[5], "Evening Shift", size_pt=8.5, bold=True, color_hex="1f2937")
    
    format_cell_text(cells_1[2], "Slot 1 (P1)", size_pt=8, bold=True, color_hex="334155")
    format_cell_text(cells_1[3], "Slot 2 (P2)", size_pt=8, bold=True, color_hex="334155")
    format_cell_text(cells_1[4], "Slot 3 (P3)", size_pt=8, bold=True, color_hex="334155")
    format_cell_text(cells_1[5], "Slot 1 (P1)", size_pt=8, bold=True, color_hex="334155")
    format_cell_text(cells_1[6], "Slot 2 (P2)", size_pt=8, bold=True, color_hex="334155")
    format_cell_text(cells_1[7], "Slot 3 (P3)", size_pt=8, bold=True, color_hex="334155")
    
    # Background colors for headers
    set_cell_background(cells_0[0], "f1f5f9")
    set_cell_background(cells_0[1], "f1f5f9")
    set_cell_background(cells_0[2], "e2e8f0")
    set_cell_background(cells_0[5], "cbd5e1")
    
    for i in range(2, 5):
        set_cell_background(cells_1[i], "f8fafc")
    for i in range(5, 8):
        set_cell_background(cells_1[i], "f1f5f9")
        
    for r in range(2):
        for c in range(8):
            set_cell_margins(table.rows[r].cells[c], top=40, bottom=40, left=50, right=50)
    
    # Populate main schedules
    current_day = start_date
    row_idx = 2
    while current_day <= end_date:
        day_str = current_day.strftime("%d")
        day_name = current_day.strftime("%a")
        is_sunday = current_day.weekday() == 6 # Sunday
        
        day_shifts = schedules_by_date.get(current_day, {})
        m_shift = day_shifts.get('morning')
        e_shift = day_shifts.get('evening')
        
        m_doc1 = m_shift.doctor_1.name if (m_shift and m_shift.doctor_1) else "None"
        m_doc2 = m_shift.doctor_2.name if (m_shift and m_shift.doctor_2) else "None"
        m_doc3 = m_shift.doctor_3.name if (m_shift and m_shift.doctor_3) else "None"
        
        m_docs = [m_doc1, m_doc2, m_doc3]
        m_duplicates = {x for x in m_docs if x != "None" and m_docs.count(x) > 1}
        
        if m_doc1 in m_duplicates:
            m_doc1 = f"{m_doc1} (Duplicate)"
        if m_doc2 in m_duplicates:
            m_doc2 = f"{m_doc2} (Duplicate)"
        if m_doc3 in m_duplicates:
            m_doc3 = f"{m_doc3} (Duplicate)"
            
        e_doc1 = e_shift.doctor_1.name if (e_shift and e_shift.doctor_1) else "None"
        e_doc2 = e_shift.doctor_2.name if (e_shift and e_shift.doctor_2) else "None"
        e_doc3 = e_shift.doctor_3.name if (e_shift and e_shift.doctor_3) else "None"
        
        e_docs = [e_doc1, e_doc2, e_doc3]
        e_duplicates = {x for x in e_docs if x != "None" and e_docs.count(x) > 1}
        
        if e_doc1 in e_duplicates:
            e_doc1 = f"{e_doc1} (Duplicate)"
        if e_doc2 in e_duplicates:
            e_doc2 = f"{e_doc2} (Duplicate)"
        if e_doc3 in e_duplicates:
            e_doc3 = f"{e_doc3} (Duplicate)"
            
        row = table.add_row()
        row_cells = row.cells
        
        # Populate text
        format_cell_text(row_cells[0], day_str, size_pt=9, bold=is_sunday, color_hex="1f2937")
        format_cell_text(row_cells[1], day_name, size_pt=9, bold=is_sunday, color_hex="ef4444" if day_name in ["Sat", "Sun"] else "6b7280")
        
        format_cell_text(row_cells[2], m_doc1, size_pt=9, bold=is_sunday, color_hex="1f2937" if "(Duplicate)" not in m_doc1 else "d97706")
        format_cell_text(row_cells[3], m_doc2, size_pt=9, bold=is_sunday, color_hex="1f2937" if "(Duplicate)" not in m_doc2 else "d97706")
        format_cell_text(row_cells[4], m_doc3, size_pt=9, bold=is_sunday, color_hex="1f2937" if "(Duplicate)" not in m_doc3 else "d97706")
        format_cell_text(row_cells[5], e_doc1, size_pt=9, bold=is_sunday, color_hex="1f2937" if "(Duplicate)" not in e_doc1 else "d97706")
        format_cell_text(row_cells[6], e_doc2, size_pt=9, bold=is_sunday, color_hex="1f2937" if "(Duplicate)" not in e_doc2 else "d97706")
        format_cell_text(row_cells[7], e_doc3, size_pt=9, bold=is_sunday, color_hex="1f2937" if "(Duplicate)" not in e_doc3 else "d97706")
        
        # Shading
        # Highlight Date Column
        set_cell_background(row_cells[0], "f8fafc")
        
        # Alternating row background for clean look
        if row_idx % 2 == 0:
            for c in range(1, 8):
                set_cell_background(row_cells[c], "f8fafc")
                
        # Cell padding
        for c in range(8):
            set_cell_margins(row_cells[c], top=50, bottom=50, left=80, right=80)
            
        row_idx += 1
        current_day += datetime.timedelta(days=1)
        
    # Apply column widths to all cells in the table
    for r in table.rows:
        for c, w in enumerate(col_widths):
            r.cells[c].width = Inches(w)
            
    # Page break for workload stats & rankings (page 2)
    doc.add_page_break()
    
    # 2. Page 2 Title
    p_title2 = doc.add_paragraph()
    p_title2_run = p_title2.add_run("Doctor Workload Statistics & Equity Report")
    p_title2_run.font.name = "Arial"
    p_title2_run.font.size = Pt(14)
    p_title2_run.bold = True
    p_title2_run.font.color.rgb = RGBColor(30, 41, 59)
    p_title2.paragraph_format.space_after = Pt(15)
    
    # Outer side-by-side table container (invisible borders)
    wrapper_table = doc.add_table(rows=1, cols=3)
    clear_table_borders(wrapper_table)
    wrapper_cells = wrapper_table.rows[0].cells
    
    # Setup cell widths: Left stats table (4.5in), Gap (0.4in), Right rankings table (4.6in)
    wrapper_cells[0].width = Inches(4.5)
    wrapper_cells[1].width = Inches(0.4)
    wrapper_cells[2].width = Inches(4.6)
    
    # Add stats table in Left Cell
    left_cell = wrapper_cells[0]
    left_cell.paragraphs[0].text = "Workload Distribution Statistics"
    left_cell.paragraphs[0].runs[0].font.name = "Arial"
    left_cell.paragraphs[0].runs[0].font.size = Pt(11)
    left_cell.paragraphs[0].runs[0].font.bold = True
    left_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(51, 65, 85) # slate-700
    left_cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    
    stats_table = left_cell.add_table(rows=1, cols=5)
    set_table_borders(stats_table, color="e2e8f0")
    
    s_hdr = stats_table.rows[0].cells
    s_hdr[0].text = "Doctor Name"
    s_hdr[1].text = "Prio"
    s_hdr[2].text = "Morn"
    s_hdr[3].text = "Eve"
    s_hdr[4].text = "Total"
    
    s_widths = [1.5, 0.6, 0.7, 0.7, 0.8]
    for c, text in enumerate(["Doctor Name", "Prio", "Morn", "Eve", "Total"]):
        format_cell_text(s_hdr[c], text, size_pt=8.5, bold=True, color_hex="ffffff")
        set_cell_background(s_hdr[c], "475569") # slate-600 header
        set_cell_margins(s_hdr[c], top=60, bottom=60, left=80, right=80)
        
    s_row_idx = 1
    sorted_stats = sorted(stats.values(), key=lambda x: (x['priority'], x['name']))
    for item in sorted_stats:
        row = stats_table.add_row()
        row_cells = row.cells
        format_cell_text(row_cells[0], item['name'], size_pt=8.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell_text(row_cells[1], f"P{item['priority']}", size_pt=8.5)
        format_cell_text(row_cells[2], str(item['morning']), size_pt=8.5)
        format_cell_text(row_cells[3], str(item['evening']), size_pt=8.5)
        format_cell_text(row_cells[4], str(item['total']), size_pt=8.5, bold=True)
        
        # Alternating background colors
        if s_row_idx % 2 == 0:
            for c in range(5):
                set_cell_background(row_cells[c], "f8fafc")
                
        for c in range(5):
            set_cell_margins(row_cells[c], top=50, bottom=50, left=80, right=80)
        s_row_idx += 1
        
    # Force widths on left table
    for r in stats_table.rows:
        for c, w in enumerate(s_widths):
            r.cells[c].width = Inches(w)

    # Add rankings table in Right Cell
    right_cell = wrapper_cells[2]
    right_cell.paragraphs[0].text = "Equity Rankings & Leaves Status"
    right_cell.paragraphs[0].runs[0].font.name = "Arial"
    right_cell.paragraphs[0].runs[0].font.size = Pt(11)
    right_cell.paragraphs[0].runs[0].font.bold = True
    right_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(51, 65, 85) # slate-700
    right_cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    
    rank_table = right_cell.add_table(rows=1, cols=4)
    set_table_borders(rank_table, color="e2e8f0")
    
    r_hdr = rank_table.rows[0].cells
    r_widths = [0.5, 1.4, 1.6, 1.1]
    for c, text in enumerate(["Rank", "Doctor Name", "Leave Dates", "Equity Status"]):
        format_cell_text(r_hdr[c], text, size_pt=8.5, bold=True, color_hex="ffffff")
        set_cell_background(r_hdr[c], "475569") # slate-600 header
        set_cell_margins(r_hdr[c], top=60, bottom=60, left=80, right=80)
        
    r_row_idx = 1
    ranked_stats = sorted(stats.values(), key=lambda x: (x['total'], x['evening'], x['name']))
    for i, item in enumerate(ranked_stats):
        avg_shifts = sum(x['total'] for x in ranked_stats) / len(ranked_stats) if ranked_stats else 0
        diff = item['total'] - avg_shifts
        if diff > 3:
            status = "Overworked"
            status_color = "b91c1c" # red-700
        elif diff < -3:
            status = "Underutilized"
            status_color = "0369a1" # sky-700
        else:
            status = "Balanced"
            status_color = "15803d" # green-700
            
        leaves_str = ", ".join(item['leaves']) if item['leaves'] else "None"
        
        row = rank_table.add_row()
        row_cells = row.cells
        format_cell_text(row_cells[0], str(i+1), size_pt=8.5)
        format_cell_text(row_cells[1], item['name'], size_pt=8.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell_text(row_cells[2], leaves_str, size_pt=8.5)
        format_cell_text(row_cells[3], status, size_pt=8.5, bold=True, color_hex=status_color)
        
        # Alternating background colors
        if r_row_idx % 2 == 0:
            for c in range(4):
                set_cell_background(row_cells[c], "f8fafc")
                
        for c in range(4):
            set_cell_margins(row_cells[c], top=50, bottom=50, left=80, right=80)
        r_row_idx += 1
        
    # Force widths on right table
    for r in rank_table.rows:
        for c, w in enumerate(r_widths):
            r.cells[c].width = Inches(w)
            
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"schedule_{year}_{month:02d}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
